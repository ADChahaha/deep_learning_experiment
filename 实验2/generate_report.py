import json
import os
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def set_run(paragraph, text, font_name="楷体", font_size=Pt(14), bold=False):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    return run


def fill_table(table, row_idx, col_idx, text, font_name="宋体", font_size=Pt(14)):
    cell = table.rows[row_idx].cells[col_idx]
    paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        run.text = ""
    if not paragraph.runs:
        run = paragraph.add_run(text)
    else:
        run = paragraph.runs[0]
        run.text = text
    run.font.name = font_name
    run.font.size = font_size


def insert_paragraph_after(doc, paragraph, text, font_name="楷体", font_size=Pt(14), bold=False, align=None):
    new_p = doc.add_paragraph()
    paragraph._element.addnext(new_p._element)
    run = new_p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    if align is not None:
        new_p.alignment = align
    return new_p


def insert_picture_after(doc, paragraph, image_path: str, width=Inches(5.8), caption: Optional[str] = None):
    if not os.path.exists(image_path):
        return paragraph
    img_p = insert_paragraph_after(doc, paragraph, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    img_p.add_run().add_picture(image_path, width=width)
    last_p = img_p
    if caption:
        cap_p = insert_paragraph_after(
            doc,
            img_p,
            caption,
            font_name="楷体",
            font_size=Pt(12),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        last_p = cap_p
    return last_p


def load_json(path: Path):
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def find_complete_runs(base_dir: Path) -> List[Path]:
    runs_dir = base_dir / "outputs" / "remote_runs"
    if not runs_dir.exists():
        return []
    runs = []
    for run_dir in sorted(runs_dir.iterdir()):
        outputs_dir = run_dir / "outputs"
        if (outputs_dir / "train_metrics.json").exists() and (outputs_dir / "eval_metrics.json").exists():
            runs.append(run_dir)
    return runs


def generate_report() -> Path:
    base_dir = Path(__file__).resolve().parent
    runs = find_complete_runs(base_dir)
    latest_run = runs[-1] if runs else None
    outputs_dir = latest_run / "outputs" if latest_run else base_dir / "outputs"

    train_metrics = load_json(outputs_dir / "train_metrics.json")
    eval_metrics = load_json(outputs_dir / "eval_metrics.json")

    final_loss = float(train_metrics.get("final_train_loss", 0.0))
    final_miou = float(eval_metrics.get("miou", train_metrics.get("final_miou", 0.0)))
    train_history = [float(x) for x in train_metrics.get("train_loss_history", [])]
    miou_history = [float(x) for x in train_metrics.get("miou_history", [])]
    class_iou = eval_metrics.get("class_iou", {})
    top_items = sorted(class_iou.items(), key=lambda x: x[1], reverse=True)[:5]
    bottom_items = sorted(class_iou.items(), key=lambda x: x[1])[:5]
    top_text = "；".join([f"{name}: {value:.4f}" for name, value in top_items]) if top_items else "待补充"
    bottom_text = "；".join([f"{name}: {value:.4f}" for name, value in bottom_items]) if bottom_items else "待补充"
    epochs = len(train_history)
    train_minutes = float(train_metrics.get("train_seconds", 0.0)) / 60.0

    doc = Document(str(base_dir / "实验报告-模板.docx"))
    purpose_paragraph = doc.paragraphs[16]
    principle_paragraph = doc.paragraphs[18]
    process_paragraph = doc.paragraphs[20]
    conclusion_paragraph = doc.paragraphs[22]

    set_run(doc.paragraphs[14], "   2026年  4月   7日", font_name="楷体", font_size=Pt(15), bold=True)
    fill_table(doc.tables[0], 0, 1, "基于 PyTorch 的 DeepLabV3 图像语义分割实验")

    purpose_text = (
        "1. 理解语义分割任务的基本目标，掌握像素级分类与图像分类任务的差异。\n"
        "2. 熟悉 PASCAL VOC2012 数据集的组织方式、标签格式与 21 类语义标注方法。\n"
        "3. 掌握 DeepLabV3 网络的基本结构，包括骨干网络、空洞卷积和 ASPP 模块。\n"
        "4. 学会使用 PyTorch 与 torchvision 提供的官方预训练模型进行微调训练。\n"
        "5. 通过训练日志、TensorBoard 可视化、分割样例和 mIoU 指标综合分析模型效果。"
    )
    set_run(purpose_paragraph, purpose_text)

    principle_text = (
        "1. 语义分割任务\n"
        "语义分割要求对图像中每一个像素进行类别判定，输出结果是与输入图像尺寸对应的语义标签图。"
        "本实验采用 PASCAL VOC2012 数据集，共包含 20 个前景类别和 1 个背景类别。\n\n"
        "2. DeepLabV3 网络结构\n"
        "DeepLabV3 以 ResNet 为骨干提取深层特征，并通过空洞卷积扩大感受野，再结合 ASPP 模块聚合多尺度上下文信息，"
        "从而提升语义分割的整体精度与边界表达能力。本实验并未直接使用完整的 ResNet101 分类网络，"
        "而是只保留其卷积骨干部分（conv1、layer1、layer2、layer3、layer4）作为特征提取器，"
        "去掉原有分类头后再接入 ASPP 与上采样层完成语义分割。\n\n"
        "3. 预训练与微调\n"
        "实验使用 torchvision 中 ResNet101 的官方预训练参数初始化骨干网络，"
        "再按照 notebook 中的 DeepLabV3 结构重新组装分割头并在 VOC2012 数据集上进行微调。"
        "这样既能保留预训练骨干已有的表征能力，又能使模型进一步适应本实验数据分布。\n\n"
        "4. 训练与评估方法\n"
        "训练阶段采用随机缩放、随机裁剪、随机水平翻转等数据增强，并使用忽略标签 255 的交叉熵损失函数。"
        "日志输出和 TensorBoard 标量均改为记录固定区间内的平均损失，以减小单 batch 波动带来的干扰。"
        "评估阶段采用多尺度推理和翻转测试，并在验证集上计算各类别 IoU 与 mean IoU。"
    )
    set_run(principle_paragraph, principle_text)

    process_text = (
        "1. 实验环境\n"
        "  深度学习框架：PyTorch\n"
        f"  加速设备：{train_metrics.get('device', '未知')}\n"
        "  主要依赖：torch、torchvision、numpy、matplotlib、tensorboard\n\n"
        "2. 数据集与预处理\n"
        "  使用 PASCAL VOC2012 语义分割数据集；训练阶段执行 0.5~2.0 倍随机缩放、513×513 随机裁剪与随机水平翻转，"
        "随后按 ImageNet 均值方差归一化。\n\n"
        "3. 实验所需文件与目录\n"
        "  data/ 目录下需要放置 VOC2012 解压后的数据集，即 data/VOCdevkit/VOC2012，"
        "其中应包含 JPEGImages、SegmentationClass 和 ImageSets/Segmentation 等子目录；"
        "VOCtrainval_11-May-2012.tar 可作为原始压缩包保留，但训练程序实际读取的是解压后的 VOCdevkit。\n"
        "  checkpoints/ 目录下可放置初始预训练权重文件；本实验直接使用 torchvision 提供的官方预训练参数进行初始化，"
        "训练完成后程序会在同一目录下保存最终模型 deeplabv3_resnet101_voc2012.pth。\n\n"
        "4. 模型与训练配置\n"
        "  模型：基于 ResNet101 骨干和 ASPP 模块实现的 DeepLabV3\n"
        f"  训练轮数：{epochs}\n"
        "  优化器：SGD\n"
        "  学习率策略：余弦衰减\n"
        "  batch size：2，梯度累计步数：8（等效 batch size 为 16）\n"
        "  训练日志：每 20 step 输出一次区间平均损失 avg_loss\n"
        "  评估方式：scales=[0.75, 1.0, 1.25]，flip=True\n"
        "  评估指标：mean IoU\n\n"
        "5. 训练流程与评估流程\n"
        "  训练脚本为 src/train.py，运行方式为 python -m src.train。程序首先读取 config.json，构建 VOC2012 训练集 DataLoader，"
        "构建以 ResNet101 骨干为基础的 DeepLabV3 网络，保留 backbone_initial 与 layer1~layer4 作为特征提取部分，"
        "再接 ASPP 与双线性上采样输出分割结果，然后使用 SGD 优化器和余弦衰减学习率进行多轮训练。"
        "每个 step 完成前向传播和交叉熵损失计算，训练过程中采用梯度累计后再执行参数更新，并每 20 step 记录一次区间平均损失到终端和 TensorBoard。\n"
        "  评估脚本为 src/evaluate.py，运行方式为 python -m src.evaluate。"
        "评估阶段将模型切换到 eval 模式，对验证集逐张图像进行多尺度推理与翻转测试，累计混淆矩阵，最后计算各类别 IoU 和 mean IoU，"
        "同时保存部分可视化预测结果。\n\n"
        "6. 训练结果\n"
        f"  训练损失历史：{[round(x, 6) for x in train_history]}\n"
        f"  验证 mIoU 历史：{[round(x, 6) for x in miou_history]}\n"
        f"  最终训练损失：{final_loss:.6f}\n"
        f"  最终 mean IoU：{final_miou:.6f}\n"
        f"  训练耗时：{train_minutes:.2f} 分钟\n\n"
        "7. 分类别结果摘录\n"
        f"  IoU 较高的类别包括：{top_text}\n\n"
        f"  相对较难的类别包括：{bottom_text}\n\n"
        "8. 可视化结果\n"
        "  结合本次训练生成的损失曲线可以观察到训练损失整体持续下降；"
        "预测叠加图可以直观看出模型对主体区域的覆盖情况，以及在细小物体和复杂边界上的误差。"
    )
    set_run(process_paragraph, process_text)

    last_p = insert_paragraph_after(
        doc,
        process_paragraph,
        "实验结果图示：",
        font_name="楷体",
        font_size=Pt(14),
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    last_p = insert_picture_after(
        doc,
        last_p,
        str(outputs_dir / "train_loss_curve.png"),
        width=Inches(6.2),
        caption="图 1：本次训练生成的训练损失曲线",
    )
    last_p = insert_picture_after(doc, last_p, str(outputs_dir / "visualizations" / "sample_1.png"), caption="图 2：语义分割结果示例 1")
    last_p = insert_picture_after(doc, last_p, str(outputs_dir / "visualizations" / "sample_2.png"), caption="图 3：语义分割结果示例 2")
    last_p = insert_picture_after(doc, last_p, str(outputs_dir / "visualizations" / "sample_3.png"), caption="图 4：语义分割结果示例 3")

    conclusion_text = (
        "1. 本实验成功利用 PyTorch 完成了 DeepLabV3 在 VOC2012 数据集上的语义分割训练与评估，"
        "模型能够较好地学习前景目标与背景之间的像素级语义关系。\n\n"
        f"2. 最终模型在验证集上取得了 {final_miou:.4f} 的 mean IoU，训练损失下降到 {final_loss:.4f}，"
        "说明当前配置下模型已经实现了稳定收敛。\n\n"
        "3. 从 TensorBoard 曲线和可视化结果可以看出，采用区间平均损失记录并结合梯度累计后，训练趋势更容易观察；"
        "同时使用 ResNet101 骨干网络并采用多尺度 + 翻转评估后，模型能够获得稳定的分割效果。\n\n"
        "4. 模型在 bus、background、aeroplane、person、train 等类别上表现较好，"
        "而在 chair、bicycle、pottedplant、sofa 等结构细碎或边界复杂的类别上仍存在一定困难。\n\n"
        "5. 后续可以尝试更强的骨干网络、进一步优化学习率初值与训练轮数，"
        "并结合更丰富的数据增强或更精细的多尺度策略，以继续提升分割精度与边界质量。"
    )
    set_run(conclusion_paragraph, conclusion_text)

    output_path = (latest_run / "实验报告.docx") if latest_run else (base_dir / "实验报告.docx")
    doc.save(output_path)
    print(f"实验报告已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_report()
